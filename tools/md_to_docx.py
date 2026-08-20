# -*- coding: utf-8 -*-
"""Convert a markdown delivery doc to Word (.docx) with headings/tables/code blocks.

Usage: python tools/md_to_docx.py <input.md> <output.docx>
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PRIMARY = RGBColor(0x31, 0x57, 0xD5)
DARK = RGBColor(0x22, 0x2A, 0x3A)
MUTED = RGBColor(0x5A, 0x66, 0x7A)
CODE_BG = "F2F4F8"


def set_zh_font(run, name_ascii="Calibri", name_east="微软雅黑"):
    run.font.name = name_ascii
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_east)


def shade_paragraph(paragraph, fill=CODE_BG):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_runs(paragraph, text, base_bold=False, color=None, size=None):
    tokens = re.split(r"(\*\*.*?\*\*|`[^`]*`|\*[^*]+\*|\[[^\]]*\]\([^)]*\))", text)
    for token in tokens:
        if not token:
            continue
        bold = base_bold
        italic = False
        content = token
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            content = token[2:-2]
            bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            content = token[1:-1]
            run = paragraph.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(size or 10)
            run.font.color.rgb = RGBColor(0x8A, 0x2E, 0x2E)
            set_zh_font(run, name_ascii="Consolas", name_east="等线")
            continue
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            content = token[1:-1]
            italic = True
        elif token.startswith("[") and "](" in token and token.endswith(")"):
            m = re.match(r"\[([^\]]*)\]\(([^)]*)\)", token)
            if m:
                content = m.group(1)
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
        set_zh_font(run)


def make_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], row[ci] if ci < len(row) else "", size=9)
            if ri == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_zh_font(run, name_east="微软雅黑")
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "3157D5")
                tcPr.append(shd)
    return table


def convert(src: Path, out: Path) -> None:
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    i, n = 0, len(lines)
    in_code = False
    code_buf = []
    while i < n:
        line = lines[i].rstrip()
        if line.strip().startswith("```"):
            if not in_code:
                in_code, code_buf = True, []
            else:
                in_code = False
                p = doc.add_paragraph()
                shade_paragraph(p)
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                set_zh_font(run, name_ascii="Consolas", name_east="等线")
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            j = i + 2
            rows = [header]
            while j < n and lines[j].strip().startswith("|"):
                rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            make_table(doc, rows)
            doc.add_paragraph()
            i = j
            continue
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(level=1)
            add_runs(p, stripped[2:], base_bold=True, color=PRIMARY, size=18)
        elif stripped.startswith("## "):
            p = doc.add_heading(level=2)
            add_runs(p, stripped[3:], base_bold=True, color=PRIMARY, size=15)
        elif stripped.startswith("### "):
            p = doc.add_heading(level=3)
            add_runs(p, stripped[4:], base_bold=True, color=DARK, size=12.5)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            add_runs(p, stripped[2:], color=MUTED, size=9.5)
            shade_paragraph(p, "F6F8FB")
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif re.match(r"^[-*]\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s*", "", stripped))
        else:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    convert(src, out)
