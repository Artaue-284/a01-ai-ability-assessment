from __future__ import annotations

import json
import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "提交材料（序号1和2·黑白版）"
ASSETS = OUT / "assets"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

PRIMARY = "000000"
SECONDARY = "000000"
ACCENT = "000000"
SOFT = "F2F2F2"
DARK = "000000"
MUTED = "000000"
WHITE = "FFFFFF"
LIGHT_TEAL = "F2F2F2"
LIGHT_BLUE = "E7E7E7"
LIGHT_RED = "D9D9D9"

DIMENSIONS = [
    ("basic", "AI基础认知", "理解生成式AI、模型能力与边界，能够选择合适的模型类型。"),
    ("prompt", "提示词工程", "能够明确目标、组织上下文、拆解任务并给出验收约束。"),
    ("tools", "AI工具使用", "能够选择并组合文本、图像、代码、办公与数据分析工具。"),
    ("evaluation", "结果评估与优化", "能够核验事实、识别幻觉与偏差，并以评测结果驱动改进。"),
    ("collaboration", "人机协同", "能够划分人机职责、设置检查点并处理异常与升级。"),
    ("ethics", "伦理与合规", "能够落实隐私、版权、公平、授权、审计和责任边界。"),
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = int(sum(widths) * 1440)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = f"w:{edge}"
        node = tc_borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_borders.append(node)
        for key in ("val", "sz", "space", "color"):
            if key in kwargs[edge]:
                node.set(qn(f"w:{key}"), str(kwargs[edge][key]))


def set_run_font(run, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 30, PRIMARY, 0, 18),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 18, SECONDARY, 18, 8),
        ("Heading 2", 14, PRIMARY, 12, 5),
        ("Heading 3", 11, SECONDARY, 9, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.12

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(short_title)
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A01 · AI能力测评智能体  |  第 ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    run = p.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)

    core = doc.core_properties
    core.subject = "2026数字马力杯A01项目提交材料"
    core.keywords = "A01, AI能力测评, 自适应测评, 教育评价"


def add_cover(doc: Document, title: str, subtitle: str, doc_no: str, status: str = "提交版") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("A01  |  AI 能力测评智能体")
    set_run_font(r, size=11, bold=True, color=PRIMARY)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(10)
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(20)
    p.add_run(subtitle)

    doc.add_paragraph("\n")
    meta = doc.add_table(rows=5, cols=2)
    set_table_geometry(meta, [1.45, 5.05])
    metadata = [
        ("材料序号", doc_no),
        ("项目名称", "AI 能力测评智能体"),
        ("赛题方向", "人工智能教育评价、技能测评与学习路径规划"),
        ("文档状态", status),
        ("版本日期", "2026年8月17日"),
    ]
    for i, (k, v) in enumerate(metadata):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        set_cell_shading(meta.cell(i, 0), SOFT)
        for run in meta.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, bold=True, color=SECONDARY)
        for cell in meta.rows[i].cells:
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": "DDE5E8"})

    doc.add_paragraph("\n")
    note = doc.add_table(rows=1, cols=1)
    set_table_geometry(note, [6.5])
    set_cell_shading(note.cell(0, 0), LIGHT_TEAL)
    p = note.cell(0, 0).paragraphs[0]
    r = p.add_run("范围说明  ")
    set_run_font(r, bold=True, color=PRIMARY)
    r = p.add_run("本文基于当前可运行版本撰写。百宝箱双智能体已发布并通过真实接口验证；真实学员效度、教师双评和生产级身份体系仍按未完成或部分完成标注。")
    set_run_font(r, size=9.5, color=DARK)
    doc.add_page_break()


def add_section_intro(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(10)


def add_callout(doc: Document, title: str, text: str, fill=LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": PRIMARY})
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    set_run_font(r, bold=True, color=PRIMARY)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size=9.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0].height = Inches(0.38)
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "D9D9D9")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=font_size, bold=True, color=DARK)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].text = str(value)
            if i % 2 == 1:
                set_cell_shading(cells[j], SOFT)
            for p in cells[j].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=font_size, color=DARK)
            set_cell_border(cells[j], bottom={"val": "single", "sz": "4", "color": "DDE5E8"})
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_picture(doc: Document, path: Path, caption: str, width=6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=8.5, color=MUTED)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_word_flow_diagram(doc: Document) -> None:
    steps = [
        ("01", "开始测评", "身份与题量模式"),
        ("02", "分层初测", "六维与难度覆盖"),
        ("03", "自适应补测", "薄弱项与低置信度"),
        ("04", "开放/实操", "任务过程与产物"),
        ("05", "自动评分", "规则/模型与复核"),
        ("06", "报告生成", "雷达图与学习建议"),
    ]
    table = doc.add_table(rows=3, cols=11)
    widths = [0.92, 0.18, 0.92, 0.18, 0.92, 0.18, 0.92, 0.18, 0.92, 0.18, 0.92]
    set_table_geometry(table, widths)
    for i, (number, title, desc) in enumerate(steps):
        col = i * 2
        cell = table.cell(0, col)
        cell.text = number
        set_cell_shading(cell, "E7E7E7")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(1, col)
        cell.text = title
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(2, col)
        cell.text = desc
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows:
            for run in row.cells[col].paragraphs[0].runs:
                set_run_font(run, size=8.2 if row is table.rows[2] else 9.0, bold=row is not table.rows[2], color=DARK)
        for edge in ("top", "left", "bottom", "right"):
            set_cell_border(table.cell(0, col), **{edge: {"val": "single", "sz": "6", "color": "7F7F7F"}})
            set_cell_border(table.cell(1, col), **{edge: {"val": "single", "sz": "6", "color": "7F7F7F"}})
            set_cell_border(table.cell(2, col), **{edge: {"val": "single", "sz": "6", "color": "7F7F7F"}})
        if i < 5:
            arrow = table.cell(1, col + 1)
            arrow.text = "→"
            arrow.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in arrow.paragraphs[0].runs:
                set_run_font(run, size=13, bold=True, color=DARK)
    p = doc.add_paragraph("图 1  学员从开始测评到报告生成的完整流程（Word表格绘制，可编辑）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=8.5, color=DARK)


def add_word_ui_wireframe(doc: Document) -> None:
    table = doc.add_table(rows=7, cols=2)
    set_table_geometry(table, [4.55, 1.95])
    left = table.cell(0, 0)
    left.text = "AI 能力测评智能体｜答题页面"
    set_cell_shading(left, "D9D9D9")
    left.merge(table.cell(0, 1))
    sections = [
        ("进度区域", "进度 8 / 18　　■■■■□□□□□□"),
        ("题号导航", "[1] [2] [3] [4] [5] [6] [7] [8] [9] … [18]"),
        ("题目区域", "根据阶段显示选择题、开放题或实操任务组件。"),
        ("作答区域", "选择项 / 多行文本框 / 实操产物说明"),
        ("反馈区域", "提交后保留题面，显示错误/半对/全对、答案或量表及解析。"),
        ("操作区域", "提交本题　→　进入下一题"),
    ]
    side = ["首页\n创建/继续测评", "报告\n雷达图/证据/建议", "成长档案\n历次记录与趋势", "教学管理\n题库/复核/班级看板", "", ""]
    for i, ((label, content), side_text) in enumerate(zip(sections, side), start=1):
        cell = table.cell(i, 0)
        cell.text = f"{label}\n{content}"
        if i % 2 == 0:
            set_cell_shading(cell, "F2F2F2")
        side_cell = table.cell(i, 1)
        side_cell.text = side_text
        if side_text:
            set_cell_shading(side_cell, "E7E7E7")
        for current in (cell, side_cell):
            for p in current.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for j, run in enumerate(p.runs):
                    set_run_font(run, size=8.6, bold=j == 0, color=DARK)
            set_cell_border(current, top={"val": "single", "sz": "6", "color": "7F7F7F"}, left={"val": "single", "sz": "6", "color": "7F7F7F"}, bottom={"val": "single", "sz": "6", "color": "7F7F7F"}, right={"val": "single", "sz": "6", "color": "7F7F7F"})
    p = doc.add_paragraph("图 2  用户交互线框原型（Word表格绘制，可编辑）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=8.5, color=DARK)


def add_word_architecture(doc: Document) -> None:
    rows = [
        ("交互层", "学员测评 Web", "能力报告 / 成长档案", "教师管理 / 人工复核"),
        ("↓", "↓", "↓", "↓"),
        ("接口层", "FastAPI REST API", "鉴权与输入校验", "CSV / JSON 导入导出"),
        ("↓", "↓", "↓", "↓"),
        ("领域服务层", "自适应测评引擎", "评分适配器", "报告与统计服务"),
        ("↓", "↓", "↓", "↓"),
        ("数据层", "SQLite 当前存储", "题库 JSON 种子", "版本与审计记录"),
        ("外部模型", "Responses API 兼容适配器", "百宝箱待 Token 接入", "超时降级 / 人工复核"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1.05, 1.82, 1.82, 1.81])
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i in (0, 2, 4, 6, 7):
                set_cell_shading(cell, "E7E7E7" if i != 7 else "D9D9D9")
                set_cell_border(cell, top={"val": "single", "sz": "6", "color": "7F7F7F"}, left={"val": "single", "sz": "6", "color": "7F7F7F"}, bottom={"val": "single", "sz": "6", "color": "7F7F7F"}, right={"val": "single", "sz": "6", "color": "7F7F7F"})
            else:
                set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=8.5 if i not in (1, 3, 5) else 11, bold=i in (0, 2, 4, 6, 7), color=DARK)
    p = doc.add_paragraph("图 1  当前系统分层架构与百宝箱接入位置（Word表格绘制，可编辑）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=8.5, color=DARK)


def add_word_data_model(doc: Document) -> None:
    rows = [
        ["users\n学员/班级", "1 → N", "tests\n测评会话/引擎状态", "1 → N", "answers\n作答/得分/用时"],
        ["", "", "↓ 0..1", "", "↓"],
        ["test_feedback\n体验与歧义反馈", "", "question_items\n题目当前版本", "1 → N", "question_versions\n历史版本/操作"],
        ["", "", "题目ID → answers", "", "human_reviews / review_resolutions\n双评与争议裁决"],
    ]
    table = doc.add_table(rows=4, cols=5)
    set_table_geometry(table, [1.42, 0.55, 1.65, 0.55, 2.33])
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if value and "→" not in value and "↓" not in value:
                set_cell_shading(cell, "E7E7E7")
                set_cell_border(cell, top={"val": "single", "sz": "6", "color": "7F7F7F"}, left={"val": "single", "sz": "6", "color": "7F7F7F"}, bottom={"val": "single", "sz": "6", "color": "7F7F7F"}, right={"val": "single", "sz": "6", "color": "7F7F7F"})
            else:
                set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"})
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=8.0, bold=bool(value and "→" not in value and "↓" not in value), color=DARK)
    p = doc.add_paragraph("图 2  核心数据实体与追溯关系（Word表格绘制，可编辑）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_run_font(run, size=8.5, color=DARK)


def add_code_block(doc: Document, title: str, source: str, note: str) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(source.strip("\n").splitlines(), start=1):
        if idx > 1:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, ascii_font="Consolas", east_asia="Microsoft YaHei", size=8.0, color=DARK)
    p = doc.add_paragraph(note)
    p.paragraph_format.space_after = Pt(7)
    for run in p.runs:
        run.italic = True
        set_run_font(run, size=8.8, color=MUTED)


def font(size: int, bold=False):
    path = Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc")
    if not path.exists():
        path = Path("C:/Windows/Fonts/simhei.ttf")
    return ImageFont.truetype(str(path), size)


def rounded(draw, box, fill, outline=None, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill, spacing=6):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align="center")
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((x, y), text, font=fnt, fill=fill, spacing=spacing, align="center")


def save_process_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 760), "#F7FAFA")
    d = ImageDraw.Draw(img)
    d.text((90, 45), "A01 学员测评主流程", font=font(42, True), fill="#24364B")
    steps = [
        ("01", "开始测评", "身份与题量模式"),
        ("02", "分层初测", "六维 × 难度覆盖"),
        ("03", "自适应补测", "薄弱项与低置信度"),
        ("04", "开放/实操", "任务过程与产物"),
        ("05", "自动评分", "规则/模型 + 复核"),
        ("06", "报告生成", "雷达图与学习建议"),
    ]
    colors = ["#E7F5F4", "#EAF2F8", "#E7F5F4", "#FFF4DD", "#FCEDEE", "#EAF2F8"]
    x0, y, w, h, gap = 70, 210, 245, 285, 42
    for i, (no, title, desc) in enumerate(steps):
        x = x0 + i * (w + gap)
        rounded(d, (x, y, x + w, y + h), colors[i], "#B6C7CC", 26, 3)
        d.text((x + 24, y + 22), no, font=font(28, True), fill="#007A78")
        center_text(d, (x + 18, y + 82, x + w - 18, y + 170), title, font(29, True), "#24364B")
        center_text(d, (x + 20, y + 170, x + w - 20, y + h - 28), desc, font(22), "#61707D")
        if i < len(steps) - 1:
            ax = x + w + 9
            ay = y + h // 2
            d.line((ax, ay, ax + gap - 14, ay), fill="#19A7A0", width=7)
            d.polygon([(ax + gap - 14, ay), (ax + gap - 30, ay - 11), (ax + gap - 30, ay + 11)], fill="#19A7A0")
    d.text((70, 585), "阶段内可自由选题；完成 12 题初测后解锁自适应题组；提交后保留当前题面、作答状态与解析。", font=font(24), fill="#1F2933")
    d.text((70, 635), "当前开放/实操题由本地量表初评并进入人工复核；取得百宝箱 Token 后切换模型评分适配器。", font=font(24), fill="#1F2933")
    img.save(path)


def save_ui_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 980), "#F7FAFA")
    d = ImageDraw.Draw(img)
    d.text((80, 35), "用户交互原型与页面关系", font=font(42, True), fill="#24364B")
    # Main student screen wireframe
    rounded(d, (70, 130, 1160, 890), "#FFFFFF", "#C3D0D5", 24, 4)
    d.rectangle((70, 130, 1160, 205), fill="#24364B")
    d.text((105, 150), "AI 能力测评智能体", font=font(27, True), fill="white")
    rounded(d, (100, 235, 1130, 315), "#E7F5F4", None, 18, 0)
    d.text((125, 252), "进度  8 / 18", font=font(23, True), fill="#007A78")
    d.rounded_rectangle((340, 262, 1080, 287), radius=12, fill="#DDE9EA")
    d.rounded_rectangle((340, 262, 665, 287), radius=12, fill="#19A7A0")
    # Palette
    x, y = 110, 345
    for i in range(18):
        col, row = i % 9, i // 9
        bx = x + col * 106
        by = y + row * 75
        color = "#EAF2F8" if i < 5 else "#FFF4DD" if i < 8 else "#FFFFFF"
        outline = "#5B9BD5" if i < 5 else "#E6A23C" if i < 8 else "#B6C7CC"
        rounded(d, (bx, by, bx + 78, by + 52), color, outline, 15, 3)
        center_text(d, (bx, by, bx + 78, by + 52), str(i + 1), font(20, True), "#24364B")
    rounded(d, (105, 525, 1125, 760), "#F8FAFB", "#DDE5E8", 22, 3)
    d.text((135, 555), "题目区域", font=font(24, True), fill="#24364B")
    d.text((135, 610), "根据测评阶段显示选择题、开放题或实操任务组件。", font=font(23), fill="#61707D")
    rounded(d, (135, 680, 390, 730), "#FFFFFF", "#C3D0D5", 12, 2)
    center_text(d, (135, 680, 390, 730), "作答选项 / 文本框", font(19), "#61707D")
    rounded(d, (840, 790, 1095, 845), "#007A78", None, 14, 0)
    center_text(d, (840, 790, 1095, 845), "提交本题 / 进入下一题", font(19, True), "white")
    # Side navigation
    boxes = [
        ("首页", "创建测评 / 继续作答"),
        ("报告", "雷达图 / 证据 / 建议"),
        ("成长档案", "历次记录与趋势"),
        ("教学管理", "题库 / 复核 / 班级看板"),
    ]
    for i, (title, desc) in enumerate(boxes):
        top = 150 + i * 185
        rounded(d, (1230, top, 1730, top + 135), "#FFFFFF", "#B6C7CC", 22, 3)
        d.text((1265, top + 22), title, font=font(26, True), fill="#007A78")
        d.text((1265, top + 75), desc, font=font(20), fill="#61707D")
        d.line((1160, top + 68, 1230, top + 68), fill="#19A7A0", width=5)
    d.text((85, 925), "状态色：错误＝浅红边缘；半对＝暖黄色边缘；全对＝浅蓝边缘；未作答＝中性灰；当前题＝青绿色。", font=font(22), fill="#1F2933")
    img.save(path)


def save_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1100), "#F7FAFA")
    d = ImageDraw.Draw(img)
    d.text((75, 35), "A01 技术架构（当前实现与百宝箱接入位）", font=font(42, True), fill="#24364B")
    layers = [
        ("交互层", ["学员测评 Web", "能力报告 / 成长档案", "教师管理 / 人工复核"], "#EAF2F8"),
        ("接口层", ["FastAPI REST API", "鉴权与输入校验", "CSV / JSON 导入导出"], "#E7F5F4"),
        ("领域服务层", ["自适应测评引擎", "评分适配器", "报告与统计服务"], "#FFF4DD"),
        ("数据层", ["SQLite 当前存储", "题库 JSON 种子", "版本与审计记录"], "#F3F6F7"),
    ]
    y0 = 145
    for li, (label, items, fill) in enumerate(layers):
        y = y0 + li * 205
        d.text((80, y + 58), label, font=font(27, True), fill="#24364B")
        d.line((225, y + 75, 280, y + 75), fill="#19A7A0", width=5)
        for i, text in enumerate(items):
            x = 300 + i * 390
            rounded(d, (x, y, x + 335, y + 145), fill, "#B6C7CC", 22, 3)
            center_text(d, (x + 20, y + 15, x + 315, y + 130), text, font(23, True), "#24364B")
        # vertical arrows
        if li < len(layers) - 1:
            for i in range(3):
                x = 468 + i * 390
                d.line((x, y + 150, x, y + 193), fill="#19A7A0", width=5)
                d.polygon([(x, y + 193), (x - 9, y + 177), (x + 9, y + 177)], fill="#19A7A0")
    # External model lane
    rounded(d, (1515, 350, 1750, 730), "#FCEDEE", "#D9787F", 24, 3)
    center_text(d, (1535, 375, 1730, 475), "外部模型通道", font(25, True), "#8E3740")
    center_text(d, (1535, 475, 1730, 585), "Responses API\n兼容适配器", font(22), "#61707D")
    center_text(d, (1535, 585, 1730, 690), "百宝箱\n已发布·Token启用", font(22, True), "#8E3740")
    d.line((1280, 555, 1515, 555), fill="#D9787F", width=5)
    d.polygon([(1515, 555), (1496, 544), (1496, 566)], fill="#D9787F")
    d.text((75, 1010), "设计原则：题库、会话、报告和人工复核不依赖单一模型；模型不可用时明确降级并保留审计信息。", font=font(24), fill="#1F2933")
    img.save(path)


def save_data_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 850), "#F7FAFA")
    d = ImageDraw.Draw(img)
    d.text((80, 38), "核心数据实体与追溯关系", font=font(42, True), fill="#24364B")
    nodes = {
        "users": (90, 190, "users\n学员 / 班级"),
        "tests": (430, 190, "tests\n测评会话 / 引擎状态"),
        "answers": (790, 190, "answers\n作答 / 得分 / 用时"),
        "reviews": (1170, 120, "human_reviews\n双人独立评分"),
        "resolve": (1170, 360, "review_resolutions\n争议裁决"),
        "items": (430, 560, "question_items\n题目当前版本"),
        "versions": (790, 560, "question_versions\n历史版本 / 操作"),
        "feedback": (90, 560, "test_feedback\n体验与歧义反馈"),
    }
    for key, (x, y, label) in nodes.items():
        fill = "#E7F5F4" if key in {"tests", "answers"} else "#FFFFFF"
        rounded(d, (x, y, x + 280, y + 150), fill, "#B6C7CC", 20, 3)
        center_text(d, (x + 12, y + 12, x + 268, y + 138), label, font(22, True), "#24364B")
    arrows = [
        ((370, 265), (430, 265), "1:N"),
        ((710, 265), (790, 265), "1:N"),
        ((1070, 230), (1170, 195), "1:N"),
        ((1070, 300), (1170, 435), "0:1"),
        ((230, 560), (500, 340), "0:1"),
        ((710, 635), (790, 635), "1:N"),
        ((570, 560), (900, 340), "题目ID"),
    ]
    for (x1, y1), (x2, y2), label in arrows:
        d.line((x1, y1, x2, y2), fill="#19A7A0", width=5)
        angle = math.atan2(y2 - y1, x2 - x1)
        p1 = (x2 - 18 * math.cos(angle - 0.55), y2 - 18 * math.sin(angle - 0.55))
        p2 = (x2 - 18 * math.cos(angle + 0.55), y2 - 18 * math.sin(angle + 0.55))
        d.polygon([(x2, y2), p1, p2], fill="#19A7A0")
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        d.text((mx + 5, my - 28), label, font=font(17, True), fill="#61707D")
    d.text((80, 790), "题目采用逻辑停用与版本留痕；answers(test_id, question_id) 唯一约束阻止同场重复提交。", font=font(23), fill="#1F2933")
    img.save(path)


def save_roadmap_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 800), "#F7FAFA")
    d = ImageDraw.Draw(img)
    d.text((80, 40), "产品演进路线图", font=font(42, True), fill="#24364B")
    stages = [
        ("提交前", "稳定与证据", ["真实学员预测试", "教师双评与校准", "百宝箱接入验证"]),
        ("0–3个月", "课程试点", ["题库扩容与质检", "学习资源映射", "教师班级诊断"]),
        ("3–6个月", "平台化", ["PostgreSQL / 多租户", "统一身份与 LMS", "机构级运营看板"]),
        ("6–12个月", "生态化", ["岗位 / 证书标准", "测学练评闭环", "开放能力与数据接口"]),
    ]
    colors = ["#FCEDEE", "#FFF4DD", "#E7F5F4", "#EAF2F8"]
    x0, y, w, gap = 75, 180, 385, 45
    d.line((140, 585, 1660, 585), fill="#19A7A0", width=8)
    for i, (time, title, items) in enumerate(stages):
        x = x0 + i * (w + gap)
        rounded(d, (x, y, x + w, y + 340), colors[i], "#B6C7CC", 24, 3)
        d.text((x + 28, y + 28), time, font=font(24, True), fill="#007A78")
        d.text((x + 28, y + 82), title, font=font(30, True), fill="#24364B")
        for j, item in enumerate(items):
            d.ellipse((x + 32, y + 155 + j * 58, x + 47, y + 170 + j * 58), fill="#19A7A0")
            d.text((x + 62, y + 145 + j * 58), item, font=font(21), fill="#1F2933")
        marker_x = x + w // 2
        d.ellipse((marker_x - 15, 570, marker_x + 15, 600), fill="#007A78")
    d.text((80, 670), "每一阶段均以数据治理、人工复核和版本可追溯为底座；规模增长不以牺牲测评效度为代价。", font=font(24), fill="#1F2933")
    img.save(path)


def load_question_examples() -> dict:
    from question_bank.loader import load_all_questions

    questions = load_all_questions()
    output = {}
    for key, _, _ in DIMENSIONS:
        objective = next(q for q in questions if q["dimension"] == key and q["type"] == "single_choice" and q["difficulty"] == 1)
        performance = next(q for q in questions if q["dimension"] == key and q["type"] != "single_choice")
        output[key] = (objective, performance)
    return output


def build_function_doc(examples: dict) -> Path:
    doc = Document()
    configure_document(doc, "01 详细功能设计文档")
    doc.core_properties.title = "A01 AI能力测评智能体——详细功能设计文档"
    add_cover(doc, "详细功能设计文档", "测评流程、用户交互原型与六维题例说明", "（1）")

    add_section_intro(doc, "1. 文档目的与产品定位", "本系统面向高校学员、教师/培训师及教学管理者，通过客观题、开放作答与实操任务，在较短测评时间内形成可解释、可追溯的六维 AI 能力画像，并将结果转化为个体学习建议和群体教学决策依据。")
    add_callout(doc, "设计目标", "以“测得准、解释清、可复核、能扩展”为主线：用分层与自适应减少无效题量，用证据与置信度避免单题定级，用人工复核守住开放题评分边界。")
    doc.add_heading("1.1 用户角色与核心任务", level=2)
    add_table(doc, ["角色", "核心任务", "系统输出"], [
        ["学员", "创建/继续测评；阶段内选题；完成开放与实操任务；查看解析", "综合分、L1–L5等级、六维雷达图、证据、建议与成长记录"],
        ["教师/评审", "查看班级画像；对开放作答独立评分；处理分歧", "群体短板、题目质量、评审一致性与裁决记录"],
        ["题库管理员", "新增、编辑、启停、导入导出和版本检查", "题库统计、质量预警、操作版本与备份文件"],
        ["学校/企业", "开展课程诊断、培训评估或能力认证试点", "聚合画像、课程改进线索与可对接的数据接口"],
    ], [1.05, 2.75, 2.70])

    doc.add_heading("1.2 当前交付范围", level=2)
    add_table(doc, ["状态", "范围"], [
        ["已实现", "六维96题题库、自适应选题、15/18/25题模式、题号导航、客观/开放/实操作答、AI助手对话、结构化初评、报告、成长档案、题库后台、双人复核、统计与导出。"],
        ["待外部条件", "百宝箱正式接口与 Token；开放题真实模型评分的一致性验证；至少30名真实学员与两名教师的测评效度验证。"],
        ["本次不含", "提交材料序号（3）的演示视频或正式在线测试环境。"],
    ], [1.25, 5.25])

    doc.add_page_break()
    add_section_intro(doc, "2. 六维能力模型", "六个维度覆盖“理解—表达—使用—评价—协同—治理”的完整 AI 应用链，每个维度以 L1–L5 可观察行为描述等级，并由题目、用时、难度、开放作答和实操证据共同支撑。")
    add_table(doc, ["维度", "核心定义", "主要证据"], [
        [name, desc, evidence] for (_, name, desc), evidence in zip(DIMENSIONS, [
            "概念辨析、模型选择、能力边界解释",
            "目标清晰度、上下文编排、任务拆解、验收约束",
            "工具选型、组合流程、权限与故障处理",
            "来源核验、评价量表、偏差识别、迭代策略",
            "职责划分、检查点、异常升级、成果验收",
            "数据最小化、版权、公平、授权、审计与申诉",
        ])
    ], [1.35, 3.25, 1.90], font_size=8.8)
    doc.add_heading("2.1 等级口径", level=2)
    add_table(doc, ["等级", "分数区间", "解释口径"], [
        ["L1 入门", "0–34.5", "能够识别部分常见概念，需要结构化指导和基础练习。"],
        ["L2 基础", "34.5–49.7", "能够完成常规单步任务，但复杂场景与风险控制不足。"],
        ["L3 熟练", "49.7–65.9", "能够拆解任务、组合方法并进行基本核验。"],
        ["L4 高级", "65.9–80.8", "能够设计流程、评测体系和风险控制机制。"],
        ["L5 专家", "80.8–100", "能够建立组织级标准并持续校准；须结合开放/实操与多次证据。"],
    ], [1.10, 1.25, 4.15])
    add_callout(doc, "解释边界", "L5 不以单道选择题直接判定。分数区间与 config/level_thresholds.json（100份仿真样本校准）保持一致，正式对外宣称测评效度前应使用真实样本复校。", fill=LIGHT_RED)

    doc.add_page_break()
    add_section_intro(doc, "3. 端到端测评流程", "系统以独立 UUID 维护每名学员的会话。前12题完成六维基础覆盖，之后根据维度得分和置信度生成自适应题组，末段纳入开放题和实操任务，完成后生成个人报告。")
    add_word_flow_diagram(doc)
    doc.add_heading("3.1 流程控制规则", level=2)
    add_numbers(doc, [
        "开始：学员填写姓名、班级/团队并选择15题快速、18题标准或25题深入模式。系统创建独立 user_id 与 test_id。",
        "分层初测：前12题按六个维度轮转，前6题为难度1，后6题为难度2；同一测评按题目ID去重。",
        "自适应补测：按“置信度优先、分数次优先”排序维度，在薄弱或证据不足的维度继续抽题；根据当前分数选择难度1/2/3。",
        "开放/实操：目标题量末段至少包含一题开放作答和一题实操任务，避免只凭客观题判断高阶能力。",
        "评分：客观题确定性评分；开放与实操题由结构化量表初评并标记人工复核。已接入 ant-line 真实模型评分；百宝箱评分智能体配置令牌后启用，未配置时透明降级并显式标注。",
        "反馈与巩固：提交后留在答题界面，以错误/半对/全对三种状态着色，错误题立即显示正确答案和解析；按钮切换为“进入下一题”。",
        "报告：完成全部目标题量后计算六维分数、置信度、优势/薄弱项、题型与难度分布，并形成学习建议与成长档案。",
    ])

    doc.add_heading("3.2 三种测评模式", level=2)
    add_table(doc, ["模式", "题量", "建议场景", "输出特点"], [
        ["快速测评", "15题", "课堂前测、活动体验", "保留六维覆盖与两类非客观任务，结论置信度相对保守"],
        ["标准测评", "18题", "常规课程诊断", "兼顾时间、维度覆盖与自适应补测"],
        ["深入测评", "25题", "队内测试、阶段认证试点", "获得更多高难度与薄弱维度证据，适合人工复核"],
    ], [1.2, 0.8, 2.0, 2.5])

    doc.add_page_break()
    add_section_intro(doc, "4. 用户交互原型", "交互遵循“当前任务突出、状态可见、错误即时巩固”的原则。题号面板置于进度条下方，阶段内允许自由选题；已提交题目可以回看，但不能重复计分。")
    add_word_ui_wireframe(doc)
    doc.add_heading("4.1 答题页关键交互", level=2)
    add_table(doc, ["区域", "交互设计", "反馈规则"], [
        ["进度与题号", "显示已答/目标题量；圆角矩形题号可点击；第二阶段未解锁题号不可操作", "当前题青绿；未答中性灰；错误浅红；半对暖黄；全对浅蓝"],
        ["题目组件", "选择题显示近似长度选项；开放/实操显示多行文本框和评分要点提示", "防止空答案提交；保留用户原始作答"],
        ["提交后", "页面不跳转，显示得分状态、正确答案/评分量表和解析", "按钮由“提交本题”替换为“进入下一题”"],
        ["中断续答", "浏览器保存学员ID；重新进入时恢复未完成测评", "后端数据库保存引擎状态与作答记录"],
    ], [1.05, 3.15, 2.30], font_size=8.7)

    doc.add_heading("4.2 报告与教学管理原型", level=2)
    add_bullets(doc, [
        "个人报告：综合分与L1–L5等级、六维雷达图、各维置信度、逐题证据、优势/薄弱项和可执行建议。",
        "成长档案：按同一学员ID聚合历次完成记录，呈现分数变化与阶段性趋势。",
        "班级看板：统计参与人数、完成数与六维平均分，帮助教师识别共性短板。",
        "题库管理：按ID、题干和标签检索；新增、编辑、启停、批量导入导出并保留版本。",
        "人工复核：开放作答进入待复核队列，两名评审独立评分，差异较大时由第三人/负责人裁决。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "5. 功能模块设计", "系统以学员端、教学管理端和公共支撑能力三部分构成。下表可直接作为开发验收与演示脚本的功能核对表。")
    add_table(doc, ["模块", "功能点", "验收标准"], [
        ["测评启动", "身份/班级、题量模式、继续未完成测评", "创建唯一会话；不允许题库未就绪时开始"],
        ["自适应答题", "分层随机、同场去重、难度路由、阶段解锁、题号导航", "六维覆盖；已答题可回看；不能重复提交"],
        ["即时反馈", "错误/半对/全对状态、答案/量表、解析", "提交后不离开题面；错误题可立即巩固"],
        ["报告生成", "综合分、等级、六维、置信度、证据、建议", "未完成测评不可访问报告；完成后数据可追溯"],
        ["成长档案", "历次测评列表与趋势", "同一浏览器学员ID可查看自己的完成记录"],
        ["班级看板", "参与、完成、六维均值", "仅管理密钥可访问"],
        ["题库管理", "增改启停、版本、导入导出、质量统计", "修改留痕；停用代替物理删除"],
        ["人工复核", "双评、分歧裁决、相关性/Kappa过程指标", "每名评审独立；裁决覆盖最终分"],
        ["数据导出", "作答与复核CSV、题库JSON备份", "管理端鉴权；导出字段可用于教研分析"],
        ["测试反馈", "体验评分、歧义题、操作与报告反馈", "反馈与test_id关联，便于闭环修订"],
    ], [1.15, 2.75, 2.60], font_size=8.4)

    doc.add_heading("5.1 题库结构", level=2)
    add_table(doc, ["指标", "当前值", "设计说明"], [
        ["总题量", "96题", "六维15~17题，达到赛题每维10基础+5进阶的数量要求"],
        ["题型", "84选择 + 3开放 + 3实操 + 2对话 + 2代码 + 2图像", "每维至少包含一题非客观任务；支持对话、代码与图像新题型"],
        ["难度", "难度1：30；难度2：36；难度3：30", "难度由概念识别到复杂流程设计逐级提升"],
        ["答案位置", "A/B/C/D各21题", "正确选项位置完全均衡，降低位置偏差"],
        ["题目解析", "96/96", "提交后与报告复盘均可查看"],
        ["质量字段", "维度、题型、难度、标签、层级、解析、量表、区分度", "支持分层抽题、统计分析和版本治理"],
    ], [1.15, 1.55, 3.80], font_size=8.7)

    doc.add_page_break()
    add_section_intro(doc, "6. 六维题例说明", "每个维度各展示一题基础客观题和一题开放/实操任务。基础题验证概念与常规判断，非客观题用于观察任务组织、验证、责任与风险控制等高阶表现。")
    for idx, (key, name, desc) in enumerate(DIMENSIONS, start=1):
        objective, performance = examples[key]
        doc.add_heading(f"6.{idx} {name}", level=2)
        p = doc.add_paragraph(desc)
        p.paragraph_format.space_after = Pt(5)
        objective_text = f"{objective['question']}\n" + "\n".join(
            f"{'ABCD'[i]}. {opt}" for i, opt in enumerate(objective["options"])
        )
        answer_letter = "ABCD"[objective["options"].index(objective["answer"])]
        add_table(doc, ["基础客观题", "答案与测量意图"], [[
            objective_text,
            f"答案：{answer_letter}\n解析：{objective['explanation']}\n测量意图：验证学员在该维度的基础识别与判断能力。",
        ]], [4.15, 2.35], font_size=8.5)
        add_table(doc, ["开放/实操题", "评分证据"], [[
            performance["question"],
            "；".join(performance.get("rubric", [])) + f"。\n参考解析：{performance['explanation']}",
        ]], [4.15, 2.35], font_size=8.5)
        if idx in (2, 4):
            doc.add_page_break()

    doc.add_page_break()
    add_section_intro(doc, "7. 评分、报告与复核规则", "评分链将自动反馈与正式结论分开：客观题可以确定性计分；开放与实操任务输出结构化初评，但在完成真实一致性验证前统一标记为需要人工复核。")
    add_table(doc, ["题型", "评分方式", "输出", "质量控制"], [
        ["单项选择", "答案完全匹配，得满分或0分", "正确/错误、正确答案、解析", "题库盲答、选项分布、正确率与区分度检查"],
        ["开放作答", "评分量表、关键词证据、结构与内容覆盖", "0–满分、命中点、反馈、needs_review", "双人盲评；模型—人工相关；争议裁决"],
        ["实操任务", "任务分解、工具选择、过程验证、风险控制等多项量表", "错误/半对/全对、量表覆盖与改进建议", "要求产物或操作证据；必要时人工复核"],
    ], [1.05, 2.25, 1.65, 1.55], font_size=8.2)
    doc.add_heading("7.1 维度分数与置信度", level=2)
    doc.add_paragraph("每题根据难度设置证据权重：w = 0.75 + 0.25 × difficulty。系统以50分为初始先验，随累计证据量提高可靠度，逐步向实测加权正确率收敛；单题不会直接把维度分数推至0或100。置信度由证据数量、权重与表现区分度共同估计，最高显示为0.99。")
    doc.add_heading("7.2 报告解释规则", level=2)
    add_bullets(doc, [
        "综合分为六维分数的算术平均，等级仅用于形成可理解的初步画像。",
        "优势维度默认为分数不低于75分；薄弱项取当前得分最低的两个维度。",
        "每项结论都可追溯到题目ID、题型、难度、得分和用时证据。",
        "学习建议优先指向薄弱维度，并要求完成一次针对性练习后复测。",
        "未满30名真实学员、未完成教师双评前，模拟准确率不得写成真实测评准确率。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "8. 安全、隐私与异常处理", "当前版本适合本地演示和小规模预测试。正式部署需在账号、权限、传输、数据库、审计和备份方面升级。")
    add_table(doc, ["风险点", "当前措施", "正式部署增强"], [
        ["管理功能越权", "教学管理接口要求 X-Admin-Key；密钥本地生成/环境变量配置", "统一登录、RBAC、多因素认证、密钥轮换"],
        ["模型密钥泄露", "仅从环境变量读取，不写入源码和数据库", "密钥托管、调用配额、审计与异常告警"],
        ["历史题目不可追溯", "题目停用不物理删除；编辑/导入/启停产生版本记录", "题目快照、发布批次、评分量表联合版本"],
        ["开放题误判", "本地初评明确 needs_review；模型失败自动降级", "百宝箱输出JSON校验、双评抽检、低置信度人工接管"],
        ["数据与隐私", "仅采集姓名/班级等必要信息；导出受管理密钥保护", "HTTPS、最小化、保留期限、删除/匿名化、合规告知"],
    ], [1.35, 2.65, 2.50], font_size=8.4)
    doc.add_heading("8.1 主要异常及用户提示", level=2)
    add_bullets(doc, [
        "题库未达到六维题量或题型门槛：禁止新建测评并返回503与缺失项。",
        "重复提交或切换到未解锁题号：返回409并保留当前状态。",
        "模型超时/格式错误：记录警告，降级至本地量表并进入人工复核。",
        "刷新或网络中断：后端从SQLite恢复测评状态，前端继续当前会话。",
        "报告提前访问：未完成测评时拒绝生成，防止报告与作答证据不一致。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "9. 功能验收清单", "以下验收项与赛题提交材料序号（1）直接对应，可用于队内走查。")
    add_table(doc, ["验收项", "判定标准", "当前状态"], [
        ["完整流程", "开始→自适应出题→实操任务→评分→报告生成均有对应页面与接口", "已实现"],
        ["用户交互", "题号导航、阶段解锁、三态反馈、即时解析、继续作答可用", "已实现"],
        ["六维题例", "每维≥15题，包含基础、进阶及非客观任务示例", "已实现"],
        ["报告输出", "雷达图、维度分、置信度、证据、建议、历次记录", "已实现"],
        ["教师管理", "班级看板、题库、复核、统计、导入导出受鉴权保护", "已实现"],
        ["百宝箱评分", "正式 Token、结构化输出、异常降级、人工一致性验证", "待接入/待验证"],
        ["真实效度", "≥30名学员、教师双评、Kappa/相关性/区分度分析", "待真实试测"],
    ], [1.35, 3.95, 1.20], font_size=8.7)
    add_callout(doc, "提交表述建议", "文档与答辩中可展示“当前系统完整可运行、非百宝箱功能已实现”；不得把本地量表初评描述为已验证的大模型自动评分，也不得把合成模拟结果描述为真实准确率。", fill=LIGHT_RED)

    path = OUT / "01_详细功能设计文档_黑白版.docx"
    doc.save(path)
    return path


def build_technical_doc() -> Path:
    doc = Document()
    configure_document(doc, "02 技术架构与关键技术说明")
    doc.core_properties.title = "A01 AI能力测评智能体——技术架构、核心代码与关键技术选型"
    add_cover(doc, "技术架构与关键技术说明", "技术架构图、核心代码示例、测评引擎、模型调用与数据存储", "（2）")

    add_section_intro(doc, "1. 架构目标", "架构围绕可运行、可追溯、可替换和可扩展设计。当前单机版本以 FastAPI + SQLite 支撑队内测试；自适应引擎、评分适配器、报告服务和数据存储解耦，为后续百宝箱、多模态与机构级部署保留清晰扩展点。")
    add_word_architecture(doc)
    add_table(doc, ["层次", "组件", "职责"], [
        ["交互层", "原生HTML/CSS/JavaScript", "响应式单页界面；测评、报告、成长档案、管理与复核交互"],
        ["接口层", "FastAPI + Pydantic", "路由、参数校验、状态码、管理鉴权、文件导出与静态页面服务"],
        ["领域层", "AdaptiveTestEngine / LLMScorer", "分层选题、难度路由、计分置信度、结构化评分和降级"],
        ["持久层", "sqlite3 + JSON题库", "用户、会话、答案、题目版本、人工复核、反馈及报告证据"],
        ["外部能力", "Responses API兼容通道 / 百宝箱", "开放题模型评分；当前百宝箱为待Token接入项"],
    ], [1.10, 2.20, 3.20], font_size=8.7)

    doc.add_page_break()
    add_section_intro(doc, "2. 关键技术选型", "选型优先保证竞赛交付阶段的可启动与低依赖，同时明确生产化升级路径。")
    add_table(doc, ["技术", "采用原因", "局限", "演进路径"], [
        ["FastAPI", "类型校验清晰、开发效率高、自带OpenAPI文档", "当前为单体服务", "按机构/测评/题库服务拆分或保持模块化单体"],
        ["原生Web", "无构建链、部署简单、局域网/隧道即可访问", "复杂状态管理扩展成本较高", "需要时迁移Vue/React并复用现有REST API"],
        ["SQLite", "零运维、事务可靠、适合单机演示与小样本试测", "并发写入与多租户能力有限", "PostgreSQL + 连接池 + 迁移脚本"],
        ["JSON题库种子", "便于团队审阅、版本管理和批量导入", "无法独立完成发布审批", "题库数据库化、草稿/审核/发布工作流"],
        ["规则 + 模型适配器", "客观题确定性；模型故障可降级；避免锁定单一供应商", "规则初评不等同模型准确性", "百宝箱接入、量表校准、双评与抽检"],
        ["Canvas/CSS可视化", "无外部图表依赖、加载稳定", "复杂分析图扩展有限", "按需接入ECharts并保留无障碍数据表"],
    ], [1.15, 2.10, 1.55, 1.70], font_size=8.0)
    add_callout(doc, "开源与依赖说明", "核心业务逻辑为项目自研；使用 FastAPI、Pydantic、Uvicorn 等开源基础库提供 Web 服务能力。正式提交时应在 requirements 与项目说明中保留依赖名称、版本和用途。")

    doc.add_page_break()
    add_section_intro(doc, "3. 自适应测评引擎设计", "自适应机制不是纯随机抽题，而是“维度 × 题型 × 难度”的分层随机、同场去重与低置信度补测。")
    add_table(doc, ["阶段", "选题规则", "目的"], [
        ["初测 1–6", "六个维度轮转，各抽1道难度1选择题", "快速建立六维基础证据"],
        ["初测 7–12", "六个维度轮转，各抽1道难度2选择题", "验证情境应用能力"],
        ["自适应客观题", "按置信度和得分排序维度；分数<45选难度1，45–74选难度2，≥75选难度3", "向薄弱项和证据不足维度补测"],
        ["末段任务", "至少1道开放题 + 1道实操题；按排序后的维度选择", "补充高阶、过程性和表达证据"],
        ["完成判定", "达到目标题量，且开放题和实操题均已完成", "防止只依赖选择题结束测评"],
    ], [1.25, 3.65, 1.60], font_size=8.5)
    add_code_block(doc, "3.1 核心代码示例：分阶段题组生成", '''initial_target = min(12, self.state["target_questions"] - 2)
while len(slots) < initial_target:
    index = len(slots)
    dimension = DIMENSIONS[index % len(DIMENSIONS)]
    difficulty = 1 if index < 6 else 2
    question_id = self._select_candidate(
        excluded, dimension, "single_choice", difficulty
    )
    slots.append(question_id)

ranked_dimensions = sorted(
    DIMENSIONS,
    key=lambda key: (self.confidence(key), self.state["scores"][key]),
)''', "来源：algorithm/adaptive_test.py。先完成六维覆盖，再以低置信度和低分优先生成后续题组。")
    add_code_block(doc, "3.2 核心代码示例：难度与置信度", '''def _target_difficulty(self, dimension: str) -> int:
    score = self.state["scores"][dimension]
    return 1 if score < 45 else 2 if score < 75 else 3

def confidence(self, dimension: str) -> float:
    stats = self.state["dimension_stats"][dimension]
    information = 1 - math.exp(-stats["weighted_possible"] / 2.5)
    decisiveness = 0.8 + min(0.2, abs(accuracy - 0.5) * 0.4)
    return round(min(0.99, information * decisiveness), 3)''', "置信度表示当前证据量与表现区分度，不等同于统计置信区间；真实部署前需用学员数据校准。")

    doc.add_page_break()
    add_section_intro(doc, "4. 评分引擎与大模型调用", "评分适配器把客观题、开放题/实操题和人工复核分为三条路径。外部模型仅负责生成结构化初评，最终结论仍受量表、校验与人工复核约束。")
    add_table(doc, ["路径", "执行方式", "异常处理", "审计字段"], [
        ["客观题", "题目答案精确匹配", "无模型依赖", "question_id、answer、score、elapsed_seconds"],
        ["本地初评", "关键词覆盖65% + 长度20% + 结构15%", "始终 needs_review=true", "rubric_met、keyword_evidence、feedback、model"],
        ["远程模型", "兼容 Responses API，要求JSON对象输出", "超时/解析失败降级本地量表并记录 warning", "model、score、rubric_met、feedback、needs_review"],
        ["人工复核", "两名评审独立评分；分歧时裁决", "保留各评审原始结果", "reviewer、score、comment、rubric、resolver"],
    ], [1.10, 2.40, 1.65, 1.35], font_size=8.1)
    add_code_block(doc, "4.1 核心代码示例：模型调用与降级", '''def score(self, question: dict, answer: str) -> dict:
    if self.api_key:
        try:
            return self._remote_score(question, answer)
        except Exception as exc:
            result = self._rubric_score(question, answer)
            result["warning"] = (
                f"LLM 调用失败，已降级为规则评分：{type(exc).__name__}"
            )
            return result
    return self._rubric_score(question, answer)''', "来源：llm/scorer.py。系统不会把降级结果伪装成真实模型评分。")
    add_code_block(doc, "4.2 核心代码示例：结构化请求", '''body = json.dumps({
    "model": self.model,
    "input": [{"role": "user", "content": [
        {"type": "input_text", "text": json.dumps(prompt, ensure_ascii=False)}
    ]}],
    "text": {"format": {"type": "json_object"}},
}).encode("utf-8")''', "模型输入包含题目、评分点、满分与学员作答；输出限定为结构化 JSON，便于校验和留痕。")

    doc.add_heading("4.3 百宝箱接入方案", level=2)
    add_numbers(doc, [
        "取得平台接口文档、模型标识、Token、限额、请求/响应样例与数据使用条款。",
        "在 llm/scorer.py 新增 BaibaoxiangScorer，保持统一 score(question, answer) 返回结构。",
        "使用环境变量/密钥管理服务注入 Token；禁止提交到代码仓库或写入数据库。",
        "对 score、rubric_met、feedback、needs_review 做JSON Schema校验，并限制分数范围。",
        "设置45秒超时、有限重试、调用ID和失败降级；高风险或低置信度结果进入人工复核。",
        "使用不少于30名真实学员、两名教师双评样本，计算模型—人工相关性、加权Kappa和分歧分布。",
    ])
    add_callout(doc, "当前状态", "系统状态接口返回 baibaoxiang: pending。现有 OPENAI 兼容通道是技术预留，不代表百宝箱已完成接入，也不代表开放题自动评分准确性已验证。", fill=LIGHT_RED)

    doc.add_page_break()
    add_section_intro(doc, "5. 数据存储与追溯", "SQLite 存储学员、会话、作答、题库版本、人工复核与反馈。测评引擎状态以 JSON 保存在 tests 表中，服务重启或刷新后可恢复。")
    add_word_data_model(doc)
    add_table(doc, ["实体", "主键/约束", "关键内容"], [
        ["users", "id UUID", "姓名、班级/团队、创建时间"],
        ["tests", "id UUID；user_id外键", "状态、目标题量、state_json、开始/完成时间"],
        ["answers", "自增ID；(test_id, question_id)唯一", "原始答案、维度、题型、难度、得分、用时、反馈"],
        ["question_items", "题目ID", "当前data_json、启停、版本、来源、更新时间"],
        ["question_versions", "自增ID", "题目历史JSON、动作、操作者、时间"],
        ["human_reviews", "(answer_id, reviewer)唯一", "独立评分、评语、评分量表"],
        ["review_resolutions", "answer_id唯一", "最终裁决分、裁决人、说明"],
        ["test_feedback", "test_id唯一", "体验评分、歧义题、易用性与报告反馈"],
    ], [1.35, 2.10, 3.05], font_size=8.2)
    doc.add_heading("5.1 数据一致性策略", level=2)
    add_bullets(doc, [
        "每次测评生成独立 UUID，避免多用户共享全局会话。",
        "answers 的唯一约束阻止同一测评对同一题重复计分。",
        "题目只做逻辑停用；修改、导入和启停均记录版本，历史记录可追溯。",
        "提交答案与保存引擎状态均在请求链内完成；异常返回明确HTTP状态码。",
        "正式多机构部署迁移 PostgreSQL，并增加数据库迁移、读写隔离、备份恢复与租户字段。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "6. 接口设计", "API 以测评会话为中心，公开接口只返回作答所需字段；答案、题库管理、复核和导出接口置于管理鉴权之后。")
    add_table(doc, ["方法与路径", "用途", "主要返回/控制"], [
        ["GET /api/status", "系统、题库与评分模式检查", "题库统计、ready状态、scoring_mode、baibaoxiang"],
        ["POST /api/test/start", "创建用户与测评", "user_id、test_id、首题、进度、题号面板"],
        ["GET /api/test/{id}/palette", "恢复当前题与题号状态", "题目、进度、题号状态"],
        ["POST /api/test/{id}/select/{n}", "阶段内切换题号或回看已答题", "readonly、原答案、三态、解析/量表"],
        ["POST /api/answer/submit", "提交并评分", "得分、三态、解析、进度、题号面板"],
        ["GET /api/report/{id}", "生成完成报告", "六维、置信度、建议、逐题复盘"],
        ["GET /api/users/{id}/history", "学员成长档案", "历次完成记录与维度趋势"],
        ["/api/admin/*", "题库、看板、复核、反馈与导出", "要求 X-Admin-Key"],
    ], [2.40, 2.10, 2.00], font_size=8.0)

    doc.add_heading("6.1 作答请求时序", level=2)
    add_numbers(doc, [
        "前端提交 test_id、question_id、answer 与 elapsed_seconds。",
        "API 从数据库加载 tests.state_json 并重建 AdaptiveTestEngine。",
        "非客观题调用评分适配器；客观题直接进入引擎确定性评分。",
        "引擎校验当前待答题，更新维度统计、加权分数、历史证据和题号状态。",
        "数据库保存 answers 与最新 state_json；完成时更新 tests.status。",
        "API 返回 submitted_status、解析、正确答案/量表、进度与题号面板。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "7. 安全、部署与可扩展性", "竞赛预测试和正式生产环境采用不同安全基线。当前临时隧道只适合队内测试，正式发布应使用固定域名、HTTPS、账号权限和备份。")
    add_table(doc, ["主题", "当前实现", "生产化要求"], [
        ["启动", "Windows 一键启动；Uvicorn 监听本机/局域网", "进程托管、健康检查、日志轮转、自动重启"],
        ["远程测试", "Cloudflare Quick Tunnel 临时地址", "固定域名、反向代理、HTTPS与访问控制"],
        ["身份权限", "学员随机ID；管理端单密钥", "统一身份、RBAC、班级/机构数据隔离"],
        ["数据保护", "SQLite本地文件；管理导出鉴权", "加密传输、备份、保留期限、匿名化与审计"],
        ["模型调用", "环境变量密钥；失败降级", "密钥托管、限额、内容安全、调用审计与成本监控"],
        ["扩展题型", "schema已预留code/image类型", "沙箱执行、图像上传扫描、多模态评分与证据保存"],
    ], [1.25, 2.40, 2.85], font_size=8.1)

    doc.add_heading("7.1 测试与质量门槛", level=2)
    add_bullets(doc, [
        "题库加载校验：ID唯一、六维有效、难度合法、选择题选项/答案完整、开放题量表完整。",
        "测评就绪门槛：每维总题数≥15、客观题≥10、开放或实操题≥1。",
        "单元测试覆盖题库、会话隔离、自适应题组、评分、报告、题库管理与复核流程。",
        "算法模拟只用于发现路由与阈值问题，不替代真实学员与人工评分验证。",
        "正式验收前执行浏览器完整流程、并发会话、断网恢复、接口权限和数据库备份恢复测试。",
    ])
    add_callout(doc, "可扩展结论", "新增千题题库、百宝箱评分、多模态题型或 PostgreSQL 不需要推翻现有测评流程；主要变化集中在题库内容、评分适配器、存储实现与部署设施。")

    doc.add_page_break()
    add_section_intro(doc, "8. 技术验收矩阵", "该矩阵将赛题序号（2）的提交要求映射到可检查的设计与实现。")
    add_table(doc, ["要求", "对应实现/文档位置", "状态"], [
        ["技术架构图", "图1：交互、接口、领域、数据与外部模型分层", "已完成"],
        ["测评引擎设计", "第3章：分层初测、低置信度补测、难度路由与完成判定", "已实现"],
        ["核心代码示例", "第3–4章：题组生成、置信度、模型调用与结构化请求", "已完成"],
        ["大模型调用方式", "Responses API兼容适配器、JSON输出、失败降级、百宝箱接入方案", "通用适配已实现；百宝箱待Token"],
        ["数据存储方案", "图2与第5章：SQLite实体、约束、版本与迁移路径", "已实现"],
        ["安全与部署", "第7章：鉴权、密钥、隧道边界与生产化要求", "预测试可用；生产化待实施"],
    ], [1.40, 3.95, 1.15], font_size=8.5)

    path = OUT / "02_技术架构与关键技术说明_黑白版.docx"
    doc.save(path)
    return path


def build_future_doc() -> Path:
    doc = Document()
    configure_document(doc, "04 未来发展文档")
    doc.core_properties.title = "A01 AI能力测评智能体——未来发展文档"
    add_cover(doc, "未来发展文档", "产品迭代、市场推广与教育生态整合方案", "（4）")

    add_section_intro(doc, "1. 发展愿景与当前基线", "产品将从“可运行的AI能力诊断工具”演进为“测评—学习—练习—复测—认证”的能力成长基础设施。当前版本已具备六维题库、自适应测评、报告、成长档案、题库管理和人工复核底座；下一阶段优先解决真实效度、百宝箱评分和规模化部署。")
    add_picture(doc, ASSETS / "05_roadmap.png", "图 1  从提交验证到教育生态平台化的演进路线", width=6.25)
    add_callout(doc, "迭代原则", "先验证测评是否有效，再扩大题量与用户规模；先建立责任、审计和人工复核，再提升自动化比例。")

    doc.add_page_break()
    add_section_intro(doc, "2. 产品迭代路线", "路线按证据成熟度分阶段推进，每阶段都设置可量化的退出条件。")
    add_table(doc, ["阶段", "重点建设", "关键交付", "退出指标"], [
        ["提交前", "真实预测试与评分校准", "30名以上学员、教师双评、低质题修订、百宝箱接入验证", "完成率、Kappa/相关性、区分度与问题闭环"],
        ["0–3个月", "课程级试点", "千题题库一期、资源推荐、教师诊断周报、复测任务", "课程使用率、报告有用度、建议完成率"],
        ["3–6个月", "机构级平台", "PostgreSQL、多租户、统一身份、班级/课程配置、运营看板", "稳定性、并发、权限审计、机构续用"],
        ["6–12个月", "生态与标准", "岗位模板、证书标准、开放API、LMS/招聘/学习档案对接", "标准复用数、生态连接数、能力提升证据"],
    ], [1.05, 1.90, 2.35, 1.20], font_size=7.9)

    doc.add_heading("2.1 千题题库建设", level=2)
    add_numbers(doc, [
        "建立维度 × 能力层级 × 题型 × 难度 × 场景的题目蓝图，先补齐稀缺格子，再扩总量。",
        "采用专家编写、AI辅助草拟、第二人盲答、领域教师审核、小样本试测五级流程。",
        "记录曝光、正确率、选项分布、平均用时、区分度和歧义反馈；以版本方式修改或停用。",
        "保持同场去重与分层随机；当题库达到千题后增加曝光控制和相似题检测，降低记忆效应。",
        "对行业版本只替换场景与知识要求，六维能力标准保持可比较。",
    ])
    add_callout(doc, "规模不等于质量", "千题题库应以蓝图覆盖与统计质量为目标，不使用模板机械复制。当前96题可作为首个校准批次。")

    doc.add_heading("2.2 评分与效度升级", level=2)
    add_table(doc, ["方向", "建设内容", "验证方法"], [
        ["百宝箱评分", "量表驱动提示、结构化JSON、失败降级、调用审计、人工接管", "模型—人工相关、加权Kappa、分维度误差"],
        ["题目参数", "根据真实作答估计难度与区分度，逐步引入IRT/CAT", "交叉验证、分层样本、阈值稳定性"],
        ["等级标准", "将L1–L5与课程目标、岗位任务和可观察产物对齐", "专家效度、外部标准关联、复测稳定性"],
        ["公平性", "监测专业、年级、性别等群体差异并限制不必要特征", "DIF分析、申诉样本复核、偏差审计"],
        ["多模态任务", "代码、图像鉴别、对话过程与真实工具调用", "沙箱验证、产物评分、人工抽检"],
    ], [1.25, 3.10, 2.15], font_size=8.2)

    doc.add_page_break()
    add_section_intro(doc, "3. 个性化学习闭环", "未来报告不只指出分数，而是把每个薄弱维度映射到课程资源、训练任务和复测节点，形成可执行的学习路径。")
    add_table(doc, ["闭环步骤", "产品能力", "数据证据"], [
        ["诊断", "自适应测评与六维画像", "维度分、置信度、题型/难度、错题与用时"],
        ["推荐", "按能力层级匹配微课、案例、提示模板和风险清单", "薄弱标签、课程目标、资源先修关系"],
        ["练习", "生成针对性开放/实操任务，允许真实AI工具协作", "过程日志、产物、人工确认点"],
        ["反馈", "量表评分、解释、改进建议和教师点评", "评分点、来源证据、模型/人工版本"],
        ["复测", "间隔1–2周抽取等价题并比较成长", "同维度变化、置信度、复测稳定性"],
        ["认证", "达到标准后生成可核验能力记录", "标准版本、测评批次、复核与申诉状态"],
    ], [1.10, 3.00, 2.40], font_size=8.4)
    doc.add_heading("3.1 推荐策略", level=2)
    add_bullets(doc, [
        "低分低置信度：先补测再推荐，避免根据不足证据过度个性化。",
        "低分高置信度：推荐基础概念、示范案例和短周期刻意练习。",
        "高分低置信度：安排不同题型或更高难度任务以确认能力稳定性。",
        "高分高置信度：推荐综合项目、同伴辅导与组织级治理任务。",
        "伦理与合规维度设置硬门槛：即使综合分高，也不得跳过必要的风险训练。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "4. 教育生态整合方案", "系统通过标准API、事件与数据字典对接学校统一身份、课程平台、学习档案、证书和企业人才系统。只交换完成业务目的所需的最小数据，并保留授权、撤回和申诉机制。")
    add_table(doc, ["生态对象", "对接方式", "输入", "输出/价值"], [
        ["统一身份/教务", "OIDC/SAML、组织与课程同步", "用户、班级、课程、角色", "免重复注册、权限与范围隔离"],
        ["LMS/教学平台", "REST API、LTI、Webhook", "课程目标、资源、作业", "诊断结果、学习建议、复测任务"],
        ["学习档案/LRS", "xAPI/标准事件", "学习活动与完成记录", "能力成长轨迹与证据链"],
        ["证书/技能标准", "标准映射与签发接口", "认证规则、标准版本", "可核验能力等级与复核状态"],
        ["企业培训/招聘", "岗位模板、批量测评、报表API", "岗位任务、能力权重", "人才画像、培训差距与岗位匹配建议"],
        ["百宝箱", "服务端API适配器", "题目、量表、作答、最小上下文", "结构化初评、反馈、对话/任务能力扩展"],
    ], [1.25, 1.45, 1.65, 2.15], font_size=7.8)
    doc.add_heading("4.1 与“AI智能·教学辅具”数据打通", level=2)
    add_numbers(doc, [
        "统一能力字典：以六维ID、L1–L5等级、题目标签和课程目标作为共同语义。",
        "统一学员与课程标识：通过学校身份系统或匿名映射ID建立最小关联。",
        "事件化交换：测评完成、建议下发、资源学习、练习提交和复测完成均形成事件。",
        "资源回流：教学辅具根据薄弱标签推送微课/案例；学习完成状态回流用于安排复测。",
        "教师闭环：班级短板进入备课与教学分析，教师调整课程后观察下一周期变化。",
        "治理闭环：标准、题库、量表、模型和推荐规则均带版本号，支持撤回和审计。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "5. 市场推广策略", "市场推广先以高校课程诊断形成可信案例，再扩展到企业培训与个人学习。核心卖点不是“自动打分”，而是可解释的能力标准、较短时间的自适应诊断和可执行的教学闭环。")
    add_table(doc, ["目标市场", "首要需求", "切入产品", "价值证明"], [
        ["高校AI通识课程", "班级差异大、教师难以分层教学", "课程前测 + 班级报告 + 复测", "完成率、教师节省时间、教学调整与能力提升"],
        ["职业院校/培训机构", "实操能力和岗位能力衔接", "行业题库 + 实操任务 + 证据报告", "岗位任务完成度、证书通过率、复测提升"],
        ["企业培训", "员工AI应用水平与治理风险", "岗位模板 + 批量测评 + 管理看板", "培训ROI、流程效率、合规问题减少"],
        ["个人学习者", "不知道短板和学习顺序", "快速测评 + 学习路径 + 阶段复测", "建议采纳率、练习完成率、成长留存"],
    ], [1.35, 2.10, 1.65, 1.40], font_size=8.0)
    doc.add_heading("5.1 推广节奏", level=2)
    add_numbers(doc, [
        "设计伙伴：选择1–2门AI通识课程，与教师共同定义目标、招募样本并完成效度验证。",
        "案例沉淀：发布匿名化试点报告，展示完成时间、评审一致性、题目修订和教学改进证据。",
        "校内复制：提供课程模板、题库包、教师培训和学期诊断服务，从单门课程扩到学院。",
        "区域合作：与产教融合基地、职业教育联盟和企业培训部门联合建立行业能力模板。",
        "生态分发：通过教学平台插件/API、证书合作和岗位画像对接获得规模化触达。",
    ])
    doc.add_heading("5.2 商业与服务模式", level=2)
    add_table(doc, ["版本", "主要能力", "建议模式"], [
        ["高校课程版", "课程前/后测、班级看板、题库与报告", "按课程/学期授权或校级服务"],
        ["机构专业版", "多租户、统一身份、行业题库、数据接口", "年度SaaS或私有化部署"],
        ["企业培训版", "岗位模板、批量测评、合规治理与培训闭环", "按席位/项目/机构授权"],
        ["个人学习版", "快速诊断、训练计划、复测与成长档案", "基础免费 + 进阶学习服务"],
    ], [1.25, 3.30, 1.95], font_size=8.4)

    doc.add_page_break()
    add_section_intro(doc, "6. 运营指标与验证体系", "指标分为测评质量、用户价值、教学价值、平台稳定和商业扩展五组，防止只追求注册量或题库数量。")
    add_table(doc, ["指标组", "核心指标", "建议目标/判定"], [
        ["测评质量", "完成率、平均时长、重测稳定性、等级一致率、加权Kappa、题目区分度", "先建立基线；样本不足时只做过程监测"],
        ["评分质量", "模型—人工相关、分维度误差、低置信度率、人工改分率", "达到团队预设阈值后再提高自动化比例"],
        ["用户价值", "报告有用度、建议点击/完成率、复测率、能力提升", "报告评分≥4/5并持续跟踪"],
        ["教学价值", "教师使用率、备课调整次数、班级短板改善", "形成可复核的课程改进案例"],
        ["平台质量", "可用性、接口错误率、P95延迟、恢复时间、备份成功率", "按试点规模设SLA并逐步提升"],
        ["增长与商业", "试点转化、机构续用、活跃课程数、生态连接数", "以真实复购和使用深度验证价值"],
    ], [1.30, 3.25, 1.95], font_size=8.1)

    doc.add_page_break()
    add_section_intro(doc, "7. 风险与治理", "教育评价属于高影响场景，系统必须把测评结果定位为学习与教学决策的辅助证据，避免未经验证地用于淘汰性决策。")
    add_table(doc, ["风险", "影响", "治理措施"], [
        ["测评效度不足", "分数不能代表真实能力", "真实样本校准、专家评审、复测、外部标准关联"],
        ["模型评分偏差", "开放题对表达风格或群体产生不公平", "量表约束、双评抽检、分群分析、申诉与人工裁决"],
        ["题库泄露/记忆", "结果被刷题影响", "千题池、曝光控制、等价题、定期换卷与异常检测"],
        ["隐私与用途漂移", "个人数据被超范围使用", "最小化、告知同意、目的限制、保留期限、删除与审计"],
        ["过度依赖单一模型", "接口故障、成本或政策变化影响系统", "评分适配器、多供应商、规则降级、人工接管"],
        ["错误用于招聘/认证", "对个人产生不当影响", "仅在明确授权和效度达标后使用；提供人工复核与申诉"],
    ], [1.35, 2.10, 3.05], font_size=8.1)
    add_heading = doc.add_heading
    add_heading("7.1 决策红线", level=2)
    add_bullets(doc, [
        "未完成真实效度验证前，不将分数作为录取、淘汰或处分的唯一依据。",
        "百宝箱或其他模型只生成辅助初评，不得绕过评分量表、日志和人工复核。",
        "不得把学员原始作答用于模型再训练，除非另行取得明确、可撤回授权。",
        "不因提高完成率而删除伦理与合规维度，也不以综合高分抵消高风险表现。",
        "所有面向外部的准确率、一致性和提升数据必须注明样本、方法、时间与适用范围。",
    ])

    doc.add_page_break()
    add_section_intro(doc, "8. 未来发展验收矩阵", "该矩阵对应提交材料序号（4）的三项要求。")
    add_table(doc, ["提交要求", "本文件对应内容", "交付结论"], [
        ["产品迭代方向", "第2–3章：题库、评分、效度、多模态与个性化学习闭环", "分阶段、有退出指标"],
        ["市场推广策略", "第5章：高校、职教、企业与个人市场，试点—案例—复制—生态", "以验证数据而非宣传口号驱动"],
        ["教育生态整合", "第4章：身份、LMS、LRS、证书、招聘与百宝箱接口", "统一能力字典、最小数据与版本治理"],
        ["创新与可扩展", "千题题库、IRT/CAT、多模态、岗位/证书标准与开放API", "建立在现有模块化架构之上"],
    ], [1.35, 4.05, 1.10], font_size=8.5)
    add_callout(doc, "近期首要行动", "最有价值的下一步不是继续堆叠界面功能，而是完成真实预测试、教师双评、百宝箱接入与稳定在线部署，以形成评分有效性和可用性的可信证据。")

    path = OUT / "04_未来发展文档.docx"
    doc.save(path)
    return path


def audit_document(path: Path) -> list[str]:
    doc = Document(path)
    issues = []
    for ti, table in enumerate(doc.tables, start=1):
        grid = table._tbl.tblGrid
        if len(grid.gridCol_lst) != len(table.columns):
            issues.append(f"table {ti}: tblGrid mismatch")
        for ri, row in enumerate(table.rows, start=1):
            if len(row.cells) != len(table.columns):
                issues.append(f"table {ti} row {ri}: cell count mismatch")
    text = "\n".join(p.text for p in doc.paragraphs)
    for marker in ("TODO", "待补充内容", "xxxx"):
        if marker in text:
            issues.append(f"placeholder found: {marker}")
    return issues


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    examples = load_question_examples()
    paths = [build_function_doc(examples), build_technical_doc()]
    report = {
        "documents": [{"path": str(path), "size": path.stat().st_size, "issues": audit_document(path)} for path in paths],
        "assets": [],
    }
    (OUT / "生成校验报告.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
